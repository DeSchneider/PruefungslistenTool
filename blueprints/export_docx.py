from docx import Document
import os
import random
import re
from models import Pruefling


def ersetze_platzhalter_sicher(paragraph, ersetzungen):
    runs = paragraph.runs
    if not runs:
        return
    full_text = ''.join(run.text for run in runs)
    changed = False
    for key, wert in ersetzungen.items():
        placeholders = [
            f'{{{{ {key} }}}}',
            f'{{{{{key}}}}}',
        ]
        for placeholder in placeholders:
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, wert or '')
                changed = True
    if changed:
        runs[0].text = full_text
        for i in range(1, len(runs)):
            runs[i].text = ''


def erstelle_und_werte_judoka_block(prueflinge):
    judoka_blocks = []
    for pruefling in prueflinge:
        if not pruefling:
            continue

        grad_current = int(pruefling.kyu_grad) if pruefling.kyu_grad and pruefling.kyu_grad.isdigit() else None
        grad_old = grad_current + 1 if grad_current is not None else None
        grad_new_str = f"{grad_current}" if grad_current is not None else ""
        grad_old_str = f"{grad_old}" if grad_old is not None else ""

        datum_pruefung_str = pruefling.datum_der_pruefung.strftime('%d.%m.%Y') if pruefling.datum_der_pruefung else ""

        bewertungen = ["", "", "", "", ""]
        if pruefling.datum_der_pruefung:
            bewertungen[:4] = [random.choice(['+', '++']) for _ in range(4)]
            bewertungen[4] = random.choice(['+', '++']) if grad_current == 1 else ""

        judoka_info = {
            "judoka_name_vorname": f"{pruefling.judoka.nachname}, {pruefling.judoka.vorname}",
            "gebdat": pruefling.judoka.geburtsdatum.strftime('%d.%m.%Y') if pruefling.judoka.geburtsdatum else "",
            "judoka_verein": pruefling.judoka.verein or "",
            "dlp": datum_pruefung_str,
            "lg": grad_old_str,
            "gn": grad_new_str,
            "b1": bewertungen[0],
            "b2": bewertungen[1],
            "b3": bewertungen[2],
            "b4": bewertungen[3],
            "b5": bewertungen[4],
        }

        judoka_blocks.append(judoka_info)

    return judoka_blocks


def generiere_und_exportiere_bericht(datei_vorlage, datei_ausgabe, pruefung, prueflinge):
    doc = Document(datei_vorlage)
    prueflinge_sortiert = sorted(prueflinge, key=lambda p: int(p.kyu_grad) if p.kyu_grad.isdigit() else 999,
                                 reverse=True)

    grad_zaehler = {}
    for p in prueflinge_sortiert:
        if p.kyu_grad and p.kyu_grad.isdigit():
            grad = int(p.kyu_grad)
            grad_zaehler[grad] = grad_zaehler.get(grad, 0) + 1

    basis_daten = {
        'ausrichter': pruefung.ausrichter or '',
        'bezirk': pruefung.bezirk or '',
        'ort_strasse': pruefung.ort_strasse or '',
        'ort_ort': pruefung.ort_ort or '',
        'datum': pruefung.datum.strftime('%d.%m.%Y') if pruefung.datum else '',
        'uhrzeit_von': pruefung.uhrzeit_von.strftime('%H:%M') if pruefung.uhrzeit_von else '',
        'uhrzeit_bis': pruefung.uhrzeit_bis.strftime('%H:%M') if pruefung.uhrzeit_bis else '',
        'anzahl_pruefer': str(pruefung.prueferanzahl or 1),
        'pruefer_name': pruefung.pruefer.name if pruefung.pruefer else '',
        'pruefer_lizenz': pruefung.pruefer.lizenz_nr if pruefung.pruefer else '',
        'fremdpruefer_name': pruefung.fremdpruefer.name if pruefung.fremdpruefer else '',
        'fremdpruefer_lizenz': pruefung.fremdpruefer.lizenz_nr if pruefung.fremdpruefer else '',
        'anzahl_graduierungen': str(len(prueflinge_sortiert)),
        'antrag_nr': str(pruefung.id),
        # Neue Platzhalter
        'gg': str(len(prueflinge_sortiert)),
        'gne': '0',
        'g1': str(grad_zaehler.get(1, 0)),
        'g2': str(grad_zaehler.get(2, 0)),
        'g3': str(grad_zaehler.get(3, 0)),
        'g4': str(grad_zaehler.get(4, 0)),
        'g5': str(grad_zaehler.get(5, 0)),
        'g6': str(grad_zaehler.get(6, 0)),
        'g7': str(grad_zaehler.get(7, 0)),
        'g8': str(grad_zaehler.get(8, 0))
    }

    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text and ('{{' in paragraph_text):
            judoka_keys = ['judoka_name_vorname', 'gebdat', 'judoka_verein', 'dlp', 'lg', 'b1', 'b2', 'b3', 'b4', 'b5',
                           'gn']
            has_judoka_placeholder = any(
                any(f"{key}{i + 1}" in paragraph_text for i in range(len(prueflinge_sortiert))) for key in judoka_keys)
            if not has_judoka_placeholder:
                ersetze_platzhalter_sicher(paragraph, basis_daten)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text
                    if paragraph_text and ('{{' in paragraph_text):
                        judoka_keys = ['judoka_name_vorname', 'gebdat', 'judoka_verein', 'dlp', 'lg', 'b1', 'b2', 'b3',
                                       'b4', 'b5', 'gn']
                        has_judoka_placeholder = any(
                            any(f"{key}{i + 1}" in paragraph_text for i in range(len(prueflinge_sortiert))) for key in
                            judoka_keys)
                        if not has_judoka_placeholder:
                            ersetze_platzhalter_sicher(paragraph, basis_daten)

    judoka_text_blocks = erstelle_und_werte_judoka_block(prueflinge_sortiert)
    max_anzahl = min(20, len(judoka_text_blocks))

    for i in range(max_anzahl):
        ersetzungen = {}
        for key in ['judoka_name_vorname', 'gebdat', 'judoka_verein', 'dlp', 'lg', 'b1', 'b2', 'b3', 'b4', 'b5', 'gn']:
            ersetzungen[f"{key}{i + 1}"] = judoka_text_blocks[i].get(key, '')
        for paragraph in doc.paragraphs:
            ersetze_platzhalter_sicher(paragraph, ersetzungen)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        ersetze_platzhalter_sicher(paragraph, ersetzungen)

    placeholder_pattern = re.compile(r'\{\{\s*[\w]+\s*\}\}')

    def ersetze_uebrige_platzhalter(paragraph):
        runs = paragraph.runs
        if not runs:
            return
        full_text = ''.join(run.text for run in runs)
        if placeholder_pattern.search(full_text):
            full_text = placeholder_pattern.sub('', full_text)
            runs[0].text = full_text
            for i in range(1, len(runs)):
                runs[i].text = ''

    for paragraph in doc.paragraphs:
        ersetze_uebrige_platzhalter(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    ersetze_uebrige_platzhalter(paragraph)

    ordner = os.path.dirname(datei_ausgabe)
    if ordner and not os.path.exists(ordner):
        os.makedirs(ordner)

    doc.save(datei_ausgabe)
